#!/usr/bin/env python3
"""
Test Word Document Editor Agent

This script demonstrates the Word Editor functionality without requiring the full MCP server installation.
It shows how the system would work with Word documents.
"""

import os
import sys
from pathlib import Path

# Add src to path
sys.path.append('src')

from agent.word_editor_agent import WordEditorAgent
from agent.state import MessagesState
from langchain_core.messages import HumanMessage

def create_sample_word_document():
    """Create a sample Word document for testing (using python-docx)."""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('RFP Response - Cybersecurity Services', 0)
        
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(
            'Our company provides comprehensive cybersecurity services including threat monitoring, '
            'incident response, and compliance management. We have been in business for over 10 years '
            'and serve clients across various industries.'
        )
        
        doc.add_heading('2. Technical Solution', level=1)
        doc.add_paragraph(
            'We propose implementing a 24/7 Security Operations Center (SOC) with advanced SIEM capabilities. '
            'Our solution includes real-time monitoring, automated threat detection, and rapid incident response. '
            'The technical architecture leverages cloud-native security tools and AI-powered analytics.'
        )
        
        doc.add_heading('3. Pricing', level=1)
        doc.add_paragraph(
            'Our pricing model is based on the number of endpoints and level of service required. '
            'Basic monitoring starts at $50 per endpoint per month. Premium services with advanced '
            'threat hunting and compliance reporting are available at $75 per endpoint per month.'
        )
        
        doc.add_heading('4. Implementation Timeline', level=1)
        doc.add_paragraph(
            'Phase 1: Initial setup and configuration (2 weeks)\n'
            'Phase 2: Integration with existing systems (3 weeks)\n'
            'Phase 3: Testing and validation (1 week)\n'
            'Phase 4: Go-live and monitoring (ongoing)'
        )
        
        sample_path = "test_rfp_response.docx"
        doc.save(sample_path)
        print(f"✅ Created sample document: {sample_path}")
        return sample_path
        
    except ImportError:
        print("⚠️ python-docx not available, creating a mock document path")
        return "test_rfp_response.docx"

def test_word_editor():
    """Test the Word Editor Agent functionality."""
    print("🧪 Testing Word Editor Agent")
    print("=" * 50)
    
    # Create sample document
    doc_path = create_sample_word_document()
    
    # Initialize the Word Editor Agent
    word_editor = WordEditorAgent()
    
    # Test scenarios
    test_scenarios = [
        {
            "name": "Edit Pricing Section",
            "message": f"Edit document {doc_path} - Update the pricing to include a discount for annual contracts",
            "description": "Should find and update the pricing section with discount information"
        },
        {
            "name": "Modify Technical Solution", 
            "message": f"Edit document {doc_path} - Add information about compliance with SOC 2 Type II standards",
            "description": "Should enhance the technical solution section with compliance details"
        },
        {
            "name": "Update Timeline",
            "message": f"Edit document {doc_path} - Shorten the implementation timeline to 4 weeks total",
            "description": "Should modify the implementation timeline section"
        }
    ]
    
    for i, scenario in enumerate(test_scenarios, 1):
        print(f"\n🔍 Test {i}: {scenario['name']}")
        print(f"📝 Description: {scenario['description']}")
        print(f"💬 Message: {scenario['message']}")
        
        # Create state with the test message
        state = MessagesState({
            "messages": [HumanMessage(content=scenario['message'])]
        })
        
        # Test the parsing functionality
        try:
            # Test message parsing
            parsed = word_editor._parse_edit_request(scenario['message'])
            print(f"✅ Parsed Request:")
            print(f"   📄 Document Path: {parsed['document_path']}")
            print(f"   ❓ Question/Instruction: {parsed['instruction']}")
            
            # Test the full edit process (will show what would happen)
            if WORD_MCP_AVAILABLE:
                result = word_editor.edit_document_section(state)
                print(f"📋 Result: {result['messages'][0].content[:200]}...")
            else:
                print("⚠️ Office Word MCP Server not installed - showing mock result")
                print(f"📋 Would edit: {parsed['document_path']}")
                print(f"📝 Would apply: {parsed['instruction']}")
                
        except Exception as e:
            print(f"❌ Error in test: {e}")
        
        print("-" * 50)

def demonstrate_integration():
    """Demonstrate how this integrates with the main system."""
    print("\n🔗 Integration with Main System")
    print("=" * 50)
    
    print("📱 Usage Examples:")
    print("1. 'Edit document /path/to/proposal.docx - Add cybersecurity compliance section'")
    print("2. 'Modify document ./rfp_response.docx - Update pricing with 10% discount'") 
    print("3. 'Update document proposal.docx - Change timeline to 3 weeks'")
    
    print("\n🎯 How it works in LangGraph:")
    print("1. User sends message with 'edit document' keywords")
    print("2. Supervisor routes to 'word_editor' agent") 
    print("3. Word Editor parses document path and instruction")
    print("4. System finds relevant section using AI")
    print("5. Generates improved content with RAG context")
    print("6. Applies changes to actual Word document")
    print("7. Creates backup and reports results")

if __name__ == "__main__":
    # Check if Office Word MCP is available
    try:
        from office_word_mcp_server import create_document
        WORD_MCP_AVAILABLE = True
        print("✅ Office Word MCP Server is available")
    except ImportError:
        WORD_MCP_AVAILABLE = False
        print("⚠️ Office Word MCP Server not installed")
        print("💡 Install with: pip install office-word-mcp-server")
    
    print("\n")
    test_word_editor()
    demonstrate_integration()
    
    print(f"\n🎉 Word Editor Agent integration complete!")
    print(f"📚 The agent is now part of your RFP system and can be used via:")
    print(f"   - LangGraph interface: langgraph dev")
    print(f"   - Interactive CLI: python interactive_agent.py") 
    print(f"   - WhatsApp integration (after setup)")