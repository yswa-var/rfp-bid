"""This module provides tools for DOCX document manipulation and search functionality.

These tools allow interaction with DOCX documents including reading, updating,
and searching content within the document structure. Tools are exposed via MCP server.
"""

import asyncio
from typing import Any, Callable, List, Optional, cast, Tuple

from docx import Document
from docx.shared import Inches
from rct_agent.docx_manager import get_docx_manager


async def index_docx(docx_path: Optional[str] = None, export_json: bool = False) -> dict[str, Any]:
    """Index or re-index a DOCX document to create structured navigation and anchor mapping.
    
    This tool parses the DOCX file and creates an index of all paragraphs with their
    anchors, breadcrumbs, styles, and heading levels. Use this when you need to 
    understand the document structure or refresh the index after external changes.
    
    Args:
        docx_path: Optional path to the DOCX file. If not provided, uses the default document.
        export_json: Whether to export the index to a JSON file (default: False)
    
    Returns:
        Dict containing index statistics and structure information
    """
    manager = get_docx_manager(docx_path)
    await manager._ensure_index_loaded()
    
    paragraphs = manager.get_all_paragraphs()
    outline = manager.get_outline()
    
    result = {
        "success": True,
        "total_paragraphs": len(paragraphs),
        "total_headings": len(outline),
        "outline": outline[:10],  # Return first 10 headings for preview
        "message": f"Successfully indexed document with {len(paragraphs)} paragraphs and {len(outline)} headings"
    }
    
    if export_json:
        output_path = "document_index.json"
        await asyncio.to_thread(manager.export_index, output_path)
        result["exported_to"] = output_path
    
    return result


async def apply_edit(anchor: List[Any], new_text: str) -> dict[str, Any]:
    """Apply an edit to a specific paragraph in the DOCX document.
    
    This tool updates the text content of a paragraph identified by its anchor position.
    The anchor is a list representing the exact location: [body, table, row, col, par].
    
    Args:
        anchor: List representing [body, table, row, col, par] position, e.g. ["body", 0, 0, 0, 5]
        new_text: New text content for the paragraph
    
    Returns:
        Dict with success status and message
    """
    manager = get_docx_manager()
    await manager._ensure_index_loaded()
    success = await asyncio.to_thread(manager.update_paragraph, anchor, new_text)
    
    if success:
        return {
            "success": True,
            "message": "Edit applied successfully",
            "anchor": anchor,
            "new_text": new_text[:100] + "..." if len(new_text) > 100 else new_text
        }
    else:
        return {
            "success": False,
            "message": "Failed to apply edit. Verify the anchor is correct.",
            "anchor": anchor
        }


async def update_toc() -> dict[str, Any]:
    """Update the Table of Contents (TOC) in the DOCX document.

    This tool regenerates the table of contents based on the current heading structure.
    It extracts all headings (Heading 1-6) and creates a hierarchical TOC.

    Returns:
        Dict with the updated TOC structure and success status
    """
    manager = get_docx_manager()
    await manager._ensure_index_loaded()
    outline = manager.get_outline()

    # Build TOC structure
    toc = {
        "title": "Table of Contents",
        "entries": []
    }

    for heading in outline:
        level = heading.get("level", 0)
        text = heading.get("text", "")
        anchor = heading.get("anchor", [])

        toc["entries"].append({
            "level": level,
            "text": text,
            "anchor": anchor,
            "indent": "  " * (level - 1)
        })

    return {
        "success": True,
        "toc": toc,
        "total_entries": len(toc["entries"]),
        "message": f"Table of Contents generated with {len(toc['entries'])} entries"
    }


def _insert_content_sync(docx_path: str, content: str, section_title: Optional[str] = None) -> dict[str, Any]:
    """Synchronous helper function to insert content into DOCX document."""
    try:
        # Load the document with python-docx for editing
        doc = Document(docx_path)

        # Convert markdown content to DOCX paragraphs
        paragraphs = _markdown_to_docx_paragraphs(content)

        # Add section title if provided
        if section_title:
            # Add heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(section_title)
            heading_run.bold = True
            heading_run.font.size = Inches(0.25)  # 18pt for heading

            # Add a blank line after heading
            doc.add_paragraph()

        # Add content paragraphs
        for para_text, para_style in paragraphs:
            p = doc.add_paragraph()
            if para_style == "bold":
                run = p.add_run(para_text)
                run.bold = True
            elif para_style == "italic":
                run = p.add_run(para_text)
                run.italic = True
            else:
                p.add_run(para_text)

        # Save the document
        doc.save(docx_path)

        return {
            "success": True,
            "message": f"Successfully inserted content{' with section title: ' + section_title if section_title else ''}",
            "content_length": len(content),
            "paragraphs_added": len(paragraphs) + (1 if section_title else 0)
        }

    except Exception as e:
        return {
            "success": False,
            "message": f"Failed to insert content: {str(e)}"
        }


async def insert_content(content: str, section_title: Optional[str] = None, style: str = "Normal") -> dict[str, Any]:
    """Insert new content at the end of the DOCX document.

    This tool adds new paragraphs to the end of the document. If a section_title is provided,
    it will add a heading with the title first, then the content.

    Args:
        content: The content to insert (can be markdown format)
        section_title: Optional heading to add before the content
        style: Style for the content paragraphs ("Normal", "Heading 1", etc.)

    Returns:
        Dict with success status and message
    """
    manager = get_docx_manager()

    # Use asyncio.to_thread to run the blocking DOCX operations in a separate thread
    result = await asyncio.to_thread(_insert_content_sync, str(manager.docx_path), content, section_title)

    if result["success"]:
        # Refresh index after successful insertion
        manager._refresh_index()

    return result


def _markdown_to_docx_paragraphs(markdown_text: str) -> List[Tuple[str, str]]:
    """Convert markdown text to list of (text, style) tuples for DOCX insertion.

    Args:
        markdown_text: Markdown formatted text

    Returns:
        List of tuples containing (text, style) for each paragraph
    """
    import re

    paragraphs = []
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle headings
        if line.startswith('# '):
            paragraphs.append((line[2:], "heading1"))
        elif line.startswith('## '):
            paragraphs.append((line[3:], "heading2"))
        elif line.startswith('### '):
            paragraphs.append((line[4:], "heading3"))
        # Handle bold text
        elif '**' in line:
            # Simple bold detection
            paragraphs.append((line.replace('**', ''), "bold"))
        elif '*' in line:
            # Simple italic detection
            paragraphs.append((line.replace('*', ''), "italic"))
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            paragraphs.append((line[2:], "bullet"))
        else:
            paragraphs.append((line, "normal"))

    return paragraphs


async def get_paragraph(anchor: List[Any]) -> Optional[dict[str, Any]]:
    """Get a paragraph from the DOCX document by its anchor.
    
    Args:
        anchor: List representing [body, table, row, col, par] position, e.g. ["body", 0, 0, 0, 5]
    
    Returns:
        Dict with paragraph information including text, style, breadcrumb, and metadata
    """
    manager = get_docx_manager()
    await manager._ensure_index_loaded()
    return manager.get_paragraph(anchor)


async def search_document(query: str, case_sensitive: bool = False) -> dict[str, Any]:
    """Search for text within the DOCX document and return matching paragraphs.
    
    Args:
        query: Text to search for in the document
        case_sensitive: Whether to match case, defaults to False for case-insensitive search
    
    Returns:
        Dict with matching paragraphs, their anchors, and metadata
    """
    manager = get_docx_manager()
    await manager._ensure_index_loaded()
    matches = manager.search(query, case_sensitive)
    
    return {
        "matches": matches,
        "count": len(matches),
        "query": query
    }


async def get_document_outline() -> dict[str, Any]:
    """Get the document outline showing all headings with their structure and metadata.
    
    Returns:
        Dict with all document headings, their levels, and hierarchical structure
    """
    manager = get_docx_manager()
    await manager._ensure_index_loaded()
    outline = manager.get_outline()
    
    return {
        "headings": outline,
        "count": len(outline)
    }


# MCP-exposed tools - primary tools for external use
TOOLS: List[Callable[..., Any]] = [
    index_docx,
    apply_edit,
    update_toc,
    insert_content,
    get_paragraph,
    search_document,
    get_document_outline,
]
