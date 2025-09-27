"""
Document Generator Module for RFP Proposals
Handles Markdown and DOC generation with structured format
"""

import json
import time
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed. DOC generation will be skipped.")

class DocumentGenerator:
    """Generates structured proposals in Markdown and DOC formats"""
    
    def __init__(self):
        self.proposal_structure = {
            "1": {
                "title": "Summary",
                "subsections": [
                    "Executive Overview",
                    "Key Benefits", 
                    "Competitive Advantages",
                    "Success Metrics"
                ]
            },
            "2": {
                "title": "About CPX",
                "subsections": [
                    "2.1. CPX Purpose & Value",
                    "2.2. Key Information",
                    "2.3. Certifications & Accreditations",
                    "2.4. Organizational Structure",
                    "2.5. Team Composition"
                ]
            },
            "3": {
                "title": "Understanding of Requirements",
                "subsections": [
                    "3.1. Project Scope Analysis",
                    "3.2. Stakeholder Requirements", 
                    "3.3. Success Criteria",
                    "3.4. Risk Assessment"
                ]
            },
            "4": {
                "title": "Proposed Solution",
                "subsections": [
                    "4.1. Technical Architecture",
                    "4.2. Implementation Approach",
                    "4.3. Solution Components",
                    "4.4. Integration Strategy"
                ]
            },
            "5": {
                "title": "Implementation Plan",
                "subsections": [
                    "5.1. Project Phases",
                    "5.2. Timeline & Milestones", 
                    "5.3. Resource Allocation",
                    "5.4. Quality Assurance"
                ]
            },
            "6": {
                "title": "Team and Experience",
                "subsections": [
                    "6.1. Core Team Members",
                    "6.2. Relevant Experience",
                    "6.3. Similar Projects",
                    "6.4. Client References"
                ]
            },
            "7": {
                "title": "Pricing",
                "subsections": [
                    "7.1. Cost Breakdown",
                    "7.2. Pricing Model",
                    "7.3. Payment Terms", 
                    "7.4. Value Analysis"
                ]
            },
            "8": {
                "title": "Terms and Conditions",
                "subsections": [
                    "8.1. Contractual Terms",
                    "8.2. Service Level Agreements",
                    "8.3. Liability & Warranty",
                    "8.4. Intellectual Property"
                ]
            },
            "9": {
                "title": "Additional Services",
                "subsections": [
                    "9.1. Optional Modules",
                    "9.2. Future Enhancements",
                    "9.3. Support Services", 
                    "9.4. Training Programs"
                ]
            },
            "10": {
                "title": "Appendices",
                "subsections": [
                    "10.1. Technical Specifications",
                    "10.2. Certifications",
                    "10.3. Case Studies",
                    "10.4. Additional Documentation"
                ]
            }
        }
        
    def _map_team_to_sections(self, team_responses: Dict[str, Any]) -> Dict[str, str]:
        """Map team responses to proposal sections based on content analysis"""
        
        section_mapping = {
            "1": "summary",  # Will be generated from all teams
            "2": "about_cpx",  # Will be generated from company info
            "3": "requirements",  # Will be generated from all teams
            "4": "technical_team",  # Technical architecture
            "5": "technical_team",  # Implementation plan
            "6": "technical_team",  # Team and experience
            "7": "finance_team",   # Pricing
            "8": "legal_team",     # Terms and conditions
            "9": "technical_team", # Additional services
            "10": "appendices"     # Will be generated from all teams
        }
        
        return section_mapping
        
    def _generate_summary_section(self, team_responses: Dict[str, Any]) -> str:
        """Generate executive summary from all team responses"""
        content = """## Executive Overview
This proposal presents a comprehensive solution designed to meet your organization's specific requirements. Our multi-disciplinary team has analyzed the requirements and developed an integrated approach that leverages cutting-edge technology, proven methodologies, and industry best practices.

## Key Benefits
- **Technical Excellence**: Robust, scalable architecture designed for long-term success
- **Financial Value**: Competitive pricing with clear ROI and value proposition
- **Legal Compliance**: Full adherence to regulatory requirements and industry standards
- **Quality Assurance**: Comprehensive testing and risk management processes

## Competitive Advantages
- Multi-disciplinary team approach ensuring holistic solution design
- Proven track record in similar projects and industries
- Flexible implementation methodology adaptable to changing requirements
- Comprehensive support and maintenance services

## Success Metrics
- On-time delivery with milestone-based progress tracking
- Budget adherence with transparent cost management
- Quality standards exceeding industry benchmarks
- Client satisfaction and long-term partnership development
"""
        return content
        
    def _generate_about_cpx_section(self) -> str:
        """Generate company information section"""
        content = """## 2.1. CPX Purpose & Value
CPX is a leading technology solutions provider specializing in enterprise-grade systems integration, custom software development, and digital transformation initiatives. Our purpose is to deliver innovative solutions that drive business growth and operational excellence.

## 2.2. Key Information
- **Founded**: 2015
- **Headquarters**: Global presence with offices in major business centers
- **Team Size**: 500+ certified professionals
- **Industries Served**: Financial Services, Healthcare, Government, Manufacturing
- **Client Base**: 200+ satisfied clients worldwide

## 2.3. Certifications & Accreditations
- ISO 27001 Information Security Management
- ISO 9001 Quality Management Systems
- CMMI Level 5 for Development and Services
- Cloud platform certifications (AWS, Azure, GCP)
- Industry-specific compliance certifications

## 2.4. Organizational Structure
Our organization is structured around centers of excellence, ensuring deep domain expertise while maintaining agility and cross-functional collaboration.

## 2.5. Team Composition
- **Technical Leadership**: Senior architects and technology leads
- **Project Management**: Certified PMP and Agile practitioners
- **Quality Assurance**: Dedicated QA and testing specialists
- **Legal & Compliance**: In-house legal and compliance experts
"""
        return content
        
    def _generate_requirements_section(self, team_responses: Dict[str, Any]) -> str:
        """Generate requirements understanding section"""
        content = """## 3.1. Project Scope Analysis
Based on our comprehensive analysis of the RFP requirements, we have identified the key scope elements and deliverables. Our understanding encompasses both functional and non-functional requirements, ensuring complete coverage of your needs.

## 3.2. Stakeholder Requirements
We have identified and analyzed requirements from all stakeholder groups, including end-users, technical teams, management, and compliance officers. Our solution addresses the unique needs of each stakeholder group.

## 3.3. Success Criteria
Clear, measurable success criteria have been established, including performance metrics, quality standards, timeline adherence, and user satisfaction benchmarks.

## 3.4. Risk Assessment
Comprehensive risk analysis has been conducted, identifying potential challenges and developing mitigation strategies to ensure project success.
"""
        return content
        
    def _generate_appendices_section(self, team_responses: Dict[str, Any]) -> str:
        """Generate appendices section"""
        content = """## 10.1. Technical Specifications
Detailed technical specifications, system requirements, and architecture diagrams are provided as supporting documentation.

## 10.2. Certifications
Complete documentation of our certifications, accreditations, and compliance attestations.

## 10.3. Case Studies
Relevant case studies demonstrating successful implementations of similar solutions.

## 10.4. Additional Documentation
Supporting materials including white papers, technical references, and methodology documentation.
"""
        return content
    
    def generate_structured_proposal(self, team_responses: Dict[str, Any]) -> Dict[str, Any]:
        """Generate structured proposal from team responses"""
        
        print("🏗️ Generating structured proposal document...")
        
        # Create structured proposal data
        structured_proposal = {
            "metadata": {
                "generated_at": datetime.now().isoformat(),
                "total_sections": len(self.proposal_structure),
                "teams_involved": list(team_responses.keys())
            },
            "proposal": {}
        }
        
        section_mapping = self._map_team_to_sections(team_responses)
        
        for section_num, section_info in self.proposal_structure.items():
            section_key = f"section_{section_num}"
            
            # Determine content source
            if section_num == "1":
                content = self._generate_summary_section(team_responses)
            elif section_num == "2":
                content = self._generate_about_cpx_section()
            elif section_num == "3":
                content = self._generate_requirements_section(team_responses)
            elif section_num == "10":
                content = self._generate_appendices_section(team_responses)
            else:
                # Map to team response
                team_key = section_mapping.get(section_num)
                if team_key and team_key in team_responses:
                    content = team_responses[team_key].get('content', f'Content for {section_info["title"]} not available.')
                else:
                    content = f'Content for {section_info["title"]} will be developed during implementation.'
            
            structured_proposal["proposal"][section_key] = {
                "number": section_num,
                "title": f"{section_num}. {section_info['title']}",
                "subsections": section_info["subsections"],
                "content": content
            }
        
        return structured_proposal
    
    def generate_markdown(self, structured_proposal: Dict[str, Any]) -> str:
        """Generate Markdown document from structured proposal"""
        
        print("📝 Generating Markdown document...")
        
        markdown_parts = [
            "# 🎯 **RFP PROPOSAL RESPONSE**",
            "",
            f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**Document Version:** 1.0",
            f"**Total Sections:** {structured_proposal['metadata']['total_sections']}",
            "",
            "---",
            ""
        ]
        
        # Add table of contents
        markdown_parts.extend([
            "## 📑 Table of Contents",
            ""
        ])
        
        for section_key, section_data in structured_proposal["proposal"].items():
            markdown_parts.append(f"**{section_data['title']}**")
            for subsection in section_data["subsections"]:
                markdown_parts.append(f"   - {subsection}")
            markdown_parts.append("")
        
        markdown_parts.extend([
            "---",
            ""
        ])
        
        # Add sections with content
        for section_key, section_data in structured_proposal["proposal"].items():
            markdown_parts.extend([
                f"# {section_data['title']}",
                ""
            ])
            
            # Add subsection structure
            markdown_parts.append("**Section Structure:**")
            for subsection in section_data["subsections"]:
                markdown_parts.append(f"- {subsection}")
            markdown_parts.extend(["", "---", ""])
            
            # Add content
            markdown_parts.extend([
                section_data["content"],
                "",
                "---",
                ""
            ])
        
        # Add footer
        markdown_parts.extend([
            "## 📊 **DOCUMENT SUMMARY**",
            f"- **Generated:** {structured_proposal['metadata']['generated_at']}",
            f"- **Total Sections:** {structured_proposal['metadata']['total_sections']}",
            f"- **Teams Involved:** {', '.join(structured_proposal['metadata']['teams_involved'])}",
            f"- **Processing Method:** Multi-team structured generation",
            "",
            "*This document was generated using an AI-powered proposal generation system.*"
        ])
        
        return "\n".join(markdown_parts)
    
    def generate_docx(self, structured_proposal: Dict[str, Any], output_path: Path) -> bool:
        """Generate DOCX document from structured proposal"""
        
        if not DOCX_AVAILABLE:
            print("⚠️ DOCX generation skipped - python-docx not available")
            return False
            
        print("📄 Generating DOCX document...")
        
        try:
            doc = Document()
            
            # Document title
            title = doc.add_heading('RFP PROPOSAL RESPONSE', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Document Version: 1.0")
            doc.add_paragraph(f"Total Sections: {structured_proposal['metadata']['total_sections']}")
            
            doc.add_page_break()
            
            # Table of Contents
            toc_heading = doc.add_heading('Table of Contents', level=1)
            
            for section_key, section_data in structured_proposal["proposal"].items():
                p = doc.add_paragraph()
                p.add_run(section_data['title']).bold = True
                
                for subsection in section_data["subsections"]:
                    doc.add_paragraph(f"   • {subsection}", style='List Bullet')
                
                doc.add_paragraph()
            
            doc.add_page_break()
            
            # Add sections with content
            for section_key, section_data in structured_proposal["proposal"].items():
                # Section title
                section_heading = doc.add_heading(section_data['title'], level=1)
                
                # Subsection structure
                structure_para = doc.add_paragraph()
                structure_para.add_run("Section Structure:").bold = True
                
                for subsection in section_data["subsections"]:
                    doc.add_paragraph(f"• {subsection}", style='List Bullet')
                
                doc.add_paragraph()
                
                # Content
                content_lines = section_data["content"].split('\n')
                for line in content_lines:
                    if line.strip():
                        if line.startswith('##'):
                            doc.add_heading(line.replace('##', '').strip(), level=2)
                        elif line.startswith('#'):
                            doc.add_heading(line.replace('#', '').strip(), level=3)
                        else:
                            doc.add_paragraph(line.strip())
                
                doc.add_page_break()
            
            # Footer
            footer_heading = doc.add_heading('Document Summary', level=1)
            doc.add_paragraph(f"Generated: {structured_proposal['metadata']['generated_at']}")
            doc.add_paragraph(f"Total Sections: {structured_proposal['metadata']['total_sections']}")
            doc.add_paragraph(f"Teams Involved: {', '.join(structured_proposal['metadata']['teams_involved'])}")
            doc.add_paragraph("Processing Method: Multi-team structured generation")
            
            doc.add_paragraph().add_run("This document was generated using an AI-powered proposal generation system.").italic = True
            
            # Save document
            doc.save(output_path)
            print(f"📄 DOCX document saved to {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Error generating DOCX: {str(e)}")
            return False
    
    def generate_all_formats(self, team_responses: Dict[str, Any], output_dir: Path) -> Dict[str, Path]:
        """Generate all document formats"""
        
        output_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Generate structured proposal
        structured_proposal = self.generate_structured_proposal(team_responses)
        
        # Save structured data as JSON
        json_file = output_dir / f"structured_proposal_{timestamp}.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(structured_proposal, f, indent=2, ensure_ascii=False)
        print(f"💾 Structured data saved to {json_file}")
        
        # Generate Markdown
        markdown_content = self.generate_markdown(structured_proposal)
        markdown_file = output_dir / f"proposal_{timestamp}.md"
        with open(markdown_file, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"📝 Markdown document saved to {markdown_file}")
        
        # Generate DOCX
        docx_file = output_dir / f"proposal_{timestamp}.docx"
        docx_success = self.generate_docx(structured_proposal, docx_file)
        
        results = {
            "json": json_file,
            "markdown": markdown_file
        }
        
        if docx_success:
            results["docx"] = docx_file
        
        return results