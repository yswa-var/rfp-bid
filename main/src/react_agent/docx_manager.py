"""DOCX Manager for reading and updating DOCX documents."""

import asyncio
import shutil
from pathlib import Path
from typing import Any, Dict, List, Optional
from docx2python import docx2python
from docx import Document
from docx.shared import Inches
from react_agent.docx_indexer import DocxIndexer


class DocxManager:
    """Manage DOCX documents with read and update capabilities."""
    
    def __init__(self, docx_path: str):
        """Initialize manager with a DOCX file path."""
        self.docx_path = Path(docx_path)
        self.indexer = DocxIndexer(str(self.docx_path))
        self.index_data: List[Dict[str, Any]] = []
        self._index_loaded = False
    
    def _refresh_index(self) -> None:
        """Refresh the internal index."""
        self.index_data = self.indexer.index()
        self._index_loaded = True
    
    async def _ensure_index_loaded(self) -> None:
        """Ensure the index is loaded, loading it asynchronously if needed."""
        if not self._index_loaded:
            self.index_data = await asyncio.to_thread(self.indexer.index)
            self._index_loaded = True
    
    def get_paragraph(self, anchor: List[Any]) -> Optional[Dict[str, Any]]:
        """Get a paragraph by its anchor.
        
        Args:
            anchor: List representing [body, table, row, col, par] position
            
        Returns:
            Dictionary with paragraph info or None if not found
        """
        # Note: This method assumes the index is already loaded
        # The async wrapper in tools.py will call _ensure_index_loaded first
        for p in self.index_data:
            if p['anchor'] == anchor:
                return p
        return None
    
    def get_outline(self) -> List[Dict[str, Any]]:
        """Get document outline (headings only).
        
        Returns:
            List of heading paragraphs with metadata
        """
        return self.indexer.get_outline()
    
    def search(self, query: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """Search for paragraphs containing text.
        
        Args:
            query: Text to search for
            case_sensitive: Whether to match case
            
        Returns:
            List of matching paragraphs
        """
        return self.indexer.find_by_text(query, case_sensitive)
    
    def update_paragraph(self, anchor: List[Any], new_text: str) -> bool:
        """Update a paragraph at the given anchor.
        
        Args:
            anchor: List representing [body, table, row, col, par] position
            new_text: New text content for the paragraph
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Load the document with python-docx for editing
            doc = Document(str(self.docx_path))
            
            # Extract the position from anchor
            if len(anchor) < 5 or anchor[0] != "body":
                return False
            
            _, table_idx, row_idx, col_idx, par_idx = anchor
            
            # For simplicity, we'll update based on paragraph index
            # Note: This is a simplified approach. In production, you'd want
            # more sophisticated mapping between docx2python and python-docx structures
            
            # Try to find the paragraph by matching text
            old_para = self.get_paragraph(anchor)
            if not old_para:
                return False
            
            old_text = old_para['text']
            
            # Search for the paragraph in the document
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == old_text:
                    # Update the paragraph
                    paragraph.text = new_text
                    
                    # Save the document
                    doc.save(str(self.docx_path))
                    
                    # Refresh index
                    self._refresh_index()
                    return True
            
            return False
            
        except Exception as e:
            print(f"Error updating paragraph: {e}")
            return False
    
    def get_all_paragraphs(self) -> List[Dict[str, Any]]:
        """Get all paragraphs with metadata.
        
        Returns:
            List of all paragraphs
        """
        return self.index_data.copy()
    
    def export_index(self, output_path: str) -> None:
        """Export the index to a JSON file.
        
        Args:
            output_path: Path to save the JSON file
        """
        self.indexer.save_index(output_path)
    
    def insert_image(self, image_path: str, width: Optional[float] = None, height: Optional[float] = None, 
                     anchor: Optional[List[Any]] = None, after_anchor: Optional[List[Any]] = None,
                     position: str = "after") -> bool:
        """Insert an image into the document.
        
        Args:
            image_path: Path to the image file
            width: Image width in inches (optional)
            height: Image height in inches (optional)
            anchor: Anchor to insert image at (optional)
            after_anchor: Insert image after this anchor (optional)
            position: Position relative to anchor ("before", "after", "replace") - default "after"
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Validate image file exists
            image_path = Path(image_path)
            if not image_path.exists():
                print(f"Image file not found: {image_path}")
                return False
            
            # Load the document with python-docx for editing
            doc = Document(str(self.docx_path))
            
            # If no anchor specified, add image at the end
            if anchor is None and after_anchor is None:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                if width and height:
                    run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
                elif width:
                    run.add_picture(str(image_path), width=Inches(width))
                elif height:
                    run.add_picture(str(image_path), height=Inches(height))
                else:
                    run.add_picture(str(image_path), width=Inches(4.0))  # Default width
                
                # Save the document
                doc.save(str(self.docx_path))
                
                # Refresh index
                self._refresh_index()
                return True
            
            # Use after_anchor if provided, otherwise use anchor
            target_anchor = after_anchor if after_anchor is not None else anchor
            
            # Find the target paragraph
            target_para = self.get_paragraph(target_anchor)
            if not target_para:
                print(f"Anchor not found: {target_anchor}")
                return False
            
            target_text = target_para['text']
            
            # Find the paragraph in the document and insert image
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip() == target_text:
                    if position == "before":
                        # Insert image before this paragraph
                        new_para = doc.paragraphs[i]._element.getparent().insert_before(
                            doc._body._element.makeelement('w:p'), 
                            doc.paragraphs[i]._element
                        )
                        new_para = paragraph._parent._element.insert_before(
                            doc._body._element.makeelement('w:p'), 
                            paragraph._element
                        )
                        # Create new paragraph object
                        new_para_obj = paragraph._parent._element.insert_before(
                            doc._body._element.makeelement('w:p'), 
                            paragraph._element
                        )
                        # Add run and image
                        run = new_para_obj.add_run()
                        if width and height:
                            run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
                        elif width:
                            run.add_picture(str(image_path), width=Inches(width))
                        elif height:
                            run.add_picture(str(image_path), height=Inches(height))
                        else:
                            run.add_picture(str(image_path), width=Inches(4.0))
                    
                    elif position == "after":
                        # Insert image after this paragraph using element API
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        
                        # Create a new paragraph element
                        new_para_element = OxmlElement('w:p')
                        
                        # Insert it right after the current paragraph's element
                        paragraph._element.addnext(new_para_element)
                        
                        # Create a Paragraph object wrapper for the new element
                        from docx.text.paragraph import Paragraph
                        new_para = Paragraph(new_para_element, paragraph._parent)
                        
                        # Add run and image to the new paragraph
                        run = new_para.add_run()
                        if width and height:
                            run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
                        elif width:
                            run.add_picture(str(image_path), width=Inches(width))
                        elif height:
                            run.add_picture(str(image_path), height=Inches(height))
                        else:
                            run.add_picture(str(image_path), width=Inches(4.0))
                    
                    elif position == "replace":
                        # Replace paragraph content with image
                        paragraph.clear()
                        run = paragraph.add_run()
                        if width and height:
                            run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
                        elif width:
                            run.add_picture(str(image_path), width=Inches(width))
                        elif height:
                            run.add_picture(str(image_path), height=Inches(height))
                        else:
                            run.add_picture(str(image_path), width=Inches(4.0))
                    
                    # Save the document
                    doc.save(str(self.docx_path))
                    
                    # Refresh index
                    self._refresh_index()
                    return True
            
            print(f"Target paragraph not found in document: {target_text}")
            return False
            
        except Exception as e:
            print(f"Error inserting image: {e}")
            return False


# Global instance (will be initialized when needed)
_docx_manager: Optional[DocxManager] = None


def get_docx_manager(docx_path: Optional[str] = None) -> DocxManager:
    """Get or create the global DOCX manager instance."""
    global _docx_manager
    
    if _docx_manager is None:
        if docx_path is None:
            # Default path - you can make this configurable
            docx_path = "/Users/yash/Documents/rfp/rfp-bid/master.docx"
        _docx_manager = DocxManager(docx_path)
    
    return _docx_manager


def reset_docx_manager() -> None:
    """Reset the global DOCX manager instance."""
    global _docx_manager
    _docx_manager = None
